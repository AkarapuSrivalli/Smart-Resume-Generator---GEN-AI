import streamlit as st
import docx

def generate_resume(name, title, email, phone, address, linkedin, github, summary, skills, software, languages, experience, education, certifications, interests, template):
    doc = docx.Document()

    if template == "Template 1":
        doc.add_heading('Resume', 0)
        doc.add_heading(name, level=1)
        doc.add_paragraph(f"Title: {title}")
        doc.add_paragraph(f"Address: {address}")
        doc.add_paragraph(f"Phone: {phone}")
        doc.add_paragraph(f"Email: {email}")
        doc.add_paragraph(f"LinkedIn: {linkedin}")
        doc.add_paragraph(f"GitHub: {github}")

        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(summary)
        
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(skills)

        doc.add_heading('Software', level=1)
        doc.add_paragraph(software)
        
        doc.add_heading('Languages', level=1)
        doc.add_paragraph(languages)

        doc.add_heading('Experience', level=1)
        for job in experience:
            doc.add_heading(job.get('title', ''), level=2)
            company_info = f"{job.get('company', 'N/A')} ({job.get('start_date', 'N/A')} - {job.get('end_date', 'N/A')})"
            doc.add_paragraph(company_info)
            doc.add_paragraph(job.get('responsibilities', ''))

        doc.add_heading('Education', level=1)
        for edu in education:
            doc.add_heading(edu.get('degree', ''), level=2)
            school_info = f"{edu.get('school', 'N/A')} ({edu.get('start_date', 'N/A')} - {edu.get('end_date', 'N/A')})"
            doc.add_paragraph(school_info)
            doc.add_paragraph(edu.get('details', ''))

        doc.add_heading('Certifications', level=1)
        doc.add_paragraph(certifications)

        doc.add_heading('Interests', level=1)
        doc.add_paragraph(interests)

    # Add more templates as needed
    doc.save('resume.docx')

st.title("SmartResume Generator")
name = st.text_input("Name")
title = st.text_input("Title")
email = st.text_input("Email")
phone = st.text_input("Phone")
address = st.text_input("Address")
linkedin = st.text_input("LinkedIn Profile")
github = st.text_input("GitHub Profile")
summary = st.text_area("Professional Summary")
skills = st.text_area("Skills")
software = st.text_area("Software")
languages = st.text_area("Languages")
experience = st.text_area("Experience (e.g., Job Title, Company, Start Date, End Date, Responsibilities)")
education = st.text_area("Education (e.g., Degree, School, Start Date, End Date, Details)")
certifications = st.text_area("Certifications")
interests = st.text_area("Interests")
template = st.selectbox("Choose a Resume Template", ["Template 1"])

if st.button("Generate Resume"):
    experience_list = [dict(zip(["title", "company", "start_date", "end_date", "responsibilities"], exp.split(", "))) for exp in experience.split("\n")]
    education_list = [dict(zip(["degree", "school", "start_date", "end_date", "details"], edu.split(", "))) for edu in education.split("\n")]
    
    if all([name, title, email, phone, address, linkedin, github, summary, skills, software, languages, experience_list, education_list, certifications, interests]):
        generate_resume(name, title, email, phone, address, linkedin, github, summary, skills, software, languages, experience_list, education_list, certifications, interests, template)
        st.success("Resume generated successfully! Download the resume from the project directory.")
    else:
        st.error("Please fill out all fields before generating the resume.")